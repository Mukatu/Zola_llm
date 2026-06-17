"""Pipeline d'ingestion RAG : load → chunk → embed → insert pgvector.

Usage programmatique (à utiliser depuis un script CLI ou les workers Dramatiq) :

    from zolaos.rag.ingest import ingest_text, ingest_file
    n = await ingest_file(
        path="data/cim10/sample.txt",
        schema="rag_health",
        tags=["country:cg", "subdomain:diagnosis"],
    )

Le pipeline est idempotent au niveau (source_uri, chunk_index) grâce à la
contrainte UNIQUE — un re-run sur le même fichier ne dupliquera pas les chunks
(l'INSERT échouera, on log et on continue).
"""

from __future__ import annotations

from pathlib import Path

from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.ext.asyncio import AsyncSession

from zolaos.core.logging import get_logger
from zolaos.db.models import RAG_MODELS
from zolaos.db.session import get_session_factory
from zolaos.rag.chunking import Chunker, get_chunker
from zolaos.rag.embeddings import EmbeddingService, get_embedding_service
from zolaos.security.pii import (
    PIIRedactionPolicy,
    redact_text,
    require_policy_for_ingest,
)

_log = get_logger("zolaos.rag.ingest")


def _load_text(path: Path) -> str:
    """Charge un fichier en texte plat. Détection par extension.

    Formats supportés :
    - .txt, .md           → lecture directe UTF-8
    - .pdf                → pypdf (requis)
    - .csv                → stdlib (texte tabulaire séparé par |)
    - .xlsx               → openpyxl (toutes les feuilles concaténées)
    - .docx               → python-docx (paragraphes + cellules de tableaux)
    - .html, .htm         → BeautifulSoup (texte visible)

    Les libs lourdes sont importées lazy pour ne pas peser au démarrage.
    """
    ext = path.suffix.lower()

    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise NotImplementedError(
                "Lecture PDF nécessite `pypdf` (pip install pypdf)."
            ) from exc
        return "\n\n".join((p.extract_text() or "") for p in PdfReader(str(path)).pages)

    if ext == ".csv":
        import csv

        rows: list[str] = []
        with path.open(encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row))
        return "\n".join(rows)

    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise NotImplementedError(
                "Lecture XLSX nécessite `openpyxl` (pip install openpyxl)."
            ) from exc
        wb = load_workbook(str(path), data_only=True, read_only=True)
        out: list[str] = []
        for sheet in wb.worksheets:
            out.append(f"## Feuille: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [
                    "" if c is None else str(c).strip()
                    for c in row
                ]
                if any(cells):
                    out.append(" | ".join(cells))
            out.append("")
        wb.close()
        return "\n".join(out)

    if ext == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise NotImplementedError(
                "Lecture DOCX nécessite `python-docx` (pip install python-docx)."
            ) from exc
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # Tableaux : on les sérialise en lignes pipe-separated
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    if ext in {".html", ".htm"}:
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise NotImplementedError(
                "Lecture HTML nécessite `beautifulsoup4` (pip install beautifulsoup4)."
            ) from exc
        raw = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(raw, "html.parser")
        # Retire scripts/styles avant extraction du texte visible
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    raise NotImplementedError(
        f"Format non supporté: {ext}. "
        f"Supportés: .txt, .md, .pdf, .csv, .xlsx, .docx, .html"
    )


async def ingest_text(
    *,
    text: str,
    source_uri: str,
    schema: str,
    tags: list[str],
    pii_policy: PIIRedactionPolicy | None,
    source_id: str | None = None,
    extra_metadata: dict | None = None,
    session: AsyncSession | None = None,
    chunker: Chunker | None = None,
    embeddings: EmbeddingService | None = None,
) -> int:
    """Ingère un blob de texte. Retourne le nombre de chunks insérés.

    Args:
        pii_policy: politique d'anonymisation. OBLIGATOIRE pour les schémas
            sensibles (rag_health, rag_legal, rag_erp). Utiliser
            `PIIRedactionPolicy.NONE` uniquement pour un corpus PUBLIC.

    `session` est optionnelle : si absente, une session interne est créée et
    commit/closée. Pour ingestion en lot dans une seule transaction, passe ta
    propre session.

    Lève `ValueError` si :
    - le schéma est inconnu ;
    - le schéma est sensible et `pii_policy is None` (garde-fou Polaris).
    """
    if schema not in RAG_MODELS:
        raise ValueError(f"Schéma RAG inconnu: {schema!r}. Connus: {list(RAG_MODELS)}")
    # Garde-fou bloquant : politique explicite obligatoire pour schémas sensibles.
    policy = require_policy_for_ingest(schema, pii_policy)

    # Anonymisation pré-ingestion (avant chunking + embedding).
    text, redaction_stats = redact_text(text, policy)
    pii_meta = {
        "pii_policy": policy.value,
        "pii_stats": redaction_stats.as_dict(),
    }

    model = RAG_MODELS[schema]
    chunker = chunker or get_chunker()
    embeddings = embeddings or get_embedding_service()

    chunks = chunker.chunk(text)
    if not chunks:
        _log.info("ingest.empty", source_uri=source_uri)
        return 0

    vectors = await embeddings.aencode([c.text for c in chunks])
    if len(vectors) != len(chunks):
        raise RuntimeError(
            f"Embedding count mismatch: {len(vectors)} vs {len(chunks)} chunks"
        )

    merged_meta = {**(extra_metadata or {}), **pii_meta}
    rows = [
        {
            "source_uri": source_uri,
            "source_id": source_id,
            "chunk_index": chunk.index,
            "content": chunk.text,
            "content_tokens": chunk.tokens,
            "embedding": vector,
            "tags": tags,
            "extra_metadata": merged_meta,
        }
        for chunk, vector in zip(chunks, vectors, strict=True)
    ]

    stmt = pg_insert(model).values(rows).on_conflict_do_nothing(
        index_elements=["source_uri", "chunk_index"]
    )

    if session is not None:
        result = await session.execute(stmt)
        await session.flush()
        inserted = result.rowcount or 0
    else:
        factory = get_session_factory()
        async with factory() as new_session:
            result = await new_session.execute(stmt)
            await new_session.commit()
            inserted = result.rowcount or 0

    _log.info(
        "ingest.done",
        source_uri=source_uri,
        schema=schema,
        chunks=len(chunks),
        inserted=inserted,
        tags=tags,
    )
    return inserted


async def ingest_file(
    *,
    path: str | Path,
    schema: str,
    tags: list[str],
    pii_policy: PIIRedactionPolicy | None,
    source_id: str | None = None,
    extra_metadata: dict | None = None,
) -> int:
    """Convenience : load + ingest_text. Source URI = path absolu.

    `pii_policy` reste obligatoire pour les schémas sensibles (cf. `ingest_text`).
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)
    text = _load_text(p)
    return await ingest_text(
        text=text,
        source_uri=str(p.resolve()),
        schema=schema,
        tags=tags,
        pii_policy=pii_policy,
        source_id=source_id or p.stem,
        extra_metadata=extra_metadata,
    )
